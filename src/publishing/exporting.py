import hashlib
import html
import io
import re

from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer

from publishing.compilation import compile_manuscript
from publishing.models import ExportRecord
from publishing.profiles import profile_for

MIMES = {
    "text": "text/plain",
    "markdown": "text/markdown",
    "html": "text/html",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}
EXTENSIONS = {"text": "txt", "markdown": "md", "html": "html", "docx": "docx", "pdf": "pdf"}


def safe_filename(value):
    stem = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip("-.")[:180]
    return stem or "manuscript"


def render_text(compiled, *, markdown=False):
    lines = [f"# {compiled.title}" if markdown else compiled.title]
    if compiled.subtitle:
        lines.append(f"## {compiled.subtitle}" if markdown else compiled.subtitle)
    if compiled.author:
        lines.append(compiled.author)
    for section in compiled.sections:
        if section.heading:
            marker = (
                "## "
                if markdown and section.entry_type in ("chapter", "volume_heading")
                else ("### " if markdown else "")
            )
            lines.extend(("", marker + section.heading, ""))
        lines.extend("\n\n".join(section.paragraphs).splitlines())
        if section.entry_type == "scene":
            lines.extend(("", "* * *", ""))
    return "\n".join(lines).strip() + "\n"


def render_html(compiled):
    sections = []
    for section in compiled.sections:
        heading = f"<h2>{html.escape(section.heading)}</h2>" if section.heading else ""
        body = "".join(
            f"<p>{html.escape(value).replace(chr(10), '<br>')}</p>" for value in section.paragraphs
        )
        sections.append(
            f'<section id="entry-{html.escape(section.entry_id)}">{heading}{body}</section>'
        )
    return (
        '<!doctype html><html><head><meta charset="utf-8">'
        '<meta name="viewport" content="width=device-width">'
        f"<title>{html.escape(compiled.title)}</title><style>"
        "body{max-width:44rem;margin:3rem auto;padding:0 1rem;"
        "font:1.1rem/1.65 Georgia,serif}"
        "h1,h2{text-wrap:balance}section{margin:3rem 0}p{overflow-wrap:anywhere}"
        "@media print{nav{display:none}section{break-inside:auto}}"
        "</style></head><body>"
        f"<h1>{html.escape(compiled.title)}</h1><p>{html.escape(compiled.subtitle)}</p>{''.join(sections)}</body></html>"
    )


def render_docx(compiled, project):
    profile = profile_for(project)
    document = Document()
    properties = document.core_properties
    properties.title = compiled.title
    properties.author = compiled.author
    properties.subject = project.description[:255]
    properties.comments = f"Compiled checksum: {compiled.checksum}"
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman" if "Serif" in profile.base_font else "Courier New"
    normal.font.size = Pt(profile.font_size)
    normal.paragraph_format.line_spacing = profile.line_spacing
    normal.paragraph_format.first_line_indent = Inches(profile.paragraph_indent_inches)
    for placement in project.artwork_placements.select_related("artwork").filter(placement="cover"):
        try:
            with placement.artwork.file.open("rb") as handle:
                document.add_picture(handle, width=Inches(5.5))
        except OSError, ValueError:
            pass
    title = document.add_heading(compiled.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if compiled.subtitle:
        paragraph = document.add_paragraph(compiled.subtitle)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if compiled.author:
        paragraph = document.add_paragraph(compiled.author)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for section in compiled.sections:
        if section.page_break_before:
            document.add_page_break()
        if section.heading:
            document.add_heading(
                section.heading,
                level=1 if section.entry_type in ("chapter", "volume_heading") else 2,
            )
        for placement in project.artwork_placements.select_related("artwork").filter(
            entry_id=section.entry_id
        ):
            try:
                with placement.artwork.file.open("rb") as handle:
                    document.add_picture(handle, width=Inches(5.5))
                if placement.caption:
                    document.add_paragraph(placement.caption)
            except OSError, ValueError:
                pass
        for value in section.paragraphs:
            document.add_paragraph(value)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(compiled, project):
    profile = profile_for(project)
    output = io.BytesIO()
    page_size = A4 if profile.page_size == "a4" else LETTER
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ManuscriptBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=profile.font_size,
        leading=profile.font_size * profile.line_spacing,
        firstLineIndent=profile.paragraph_indent_inches * inch,
        spaceAfter=profile.paragraph_spacing_points,
    )
    heading = ParagraphStyle(
        "ManuscriptHeading",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        spaceBefore=18,
        spaceAfter=18,
    )
    story = [Paragraph(html.escape(compiled.title), styles["Title"])]
    for placement in project.artwork_placements.select_related("artwork").filter(placement="cover"):
        try:
            with placement.artwork.file.open("rb") as handle:
                image_data = handle.read()
            story.extend(
                (
                    Image(
                        io.BytesIO(image_data),
                        width=5.5 * inch,
                        height=7 * inch,
                        kind="proportional",
                    ),
                    Spacer(1, 12),
                )
            )
        except OSError, ValueError:
            pass
    if compiled.subtitle:
        story.append(Paragraph(html.escape(compiled.subtitle), heading))
    for section in compiled.sections:
        if section.page_break_before:
            story.append(PageBreak())
        if section.heading:
            story.append(Paragraph(html.escape(section.heading), heading))
        for placement in project.artwork_placements.select_related("artwork").filter(
            entry_id=section.entry_id
        ):
            try:
                with placement.artwork.file.open("rb") as handle:
                    image_data = handle.read()
                story.append(
                    Image(
                        io.BytesIO(image_data),
                        width=5.5 * inch,
                        height=6 * inch,
                        kind="proportional",
                    )
                )
                if placement.caption:
                    story.append(Paragraph(html.escape(placement.caption), body))
            except OSError, ValueError:
                pass
        for value in section.paragraphs:
            story.extend((Paragraph(html.escape(value).replace("\n", "<br/>"), body), Spacer(1, 6)))

    def page_number(canvas, document):
        if profile.page_numbers:
            canvas.saveState()
            canvas.setFont("Times-Roman", 9)
            canvas.drawCentredString(page_size[0] / 2, 0.45 * inch, str(document.page))
            canvas.restoreState()

    SimpleDocTemplate(
        output,
        pagesize=page_size,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=0.75 * inch,
        title=compiled.title,
        author=compiled.author,
    ).build(story, onFirstPage=page_number, onLaterPages=page_number)
    return output.getvalue()


def generate_export(record_id):
    record = ExportRecord.objects.select_related("project__work").get(id=record_id)
    if record.status == "ready" and record.file and record.checksum:
        return record
    record.status = "generating"
    record.error_message = ""
    record.save(update_fields=("status", "error_message"))
    try:
        compiled = compile_manuscript(record.project)
        if record.export_format == "text":
            data = render_text(compiled).encode()
        elif record.export_format == "markdown":
            data = render_text(compiled, markdown=True).encode()
        elif record.export_format == "html":
            data = render_html(compiled).encode()
        elif record.export_format == "docx":
            data = render_docx(compiled, record.project)
        elif record.export_format == "pdf":
            data = render_pdf(compiled, record.project)
        else:
            raise ValueError("Unsupported export format.")
        digest = hashlib.sha256(data).hexdigest()
        extension = EXTENSIONS[record.export_format]
        requested_stem = record.filename.rsplit(".", 1)[0]
        filename = f"{safe_filename(requested_stem)}-{record.id}.{extension}"
        record.file.save(filename, ContentFile(data), save=False)
        record.filename = filename
        record.mime_type = MIMES[record.export_format]
        record.file_size = len(data)
        record.checksum = digest
        record.compiled_manuscript_checksum = compiled.checksum
        record.source_snapshot = compiled.source_snapshot
        record.warning_report = list(compiled.warnings)
        record.status = "ready"
        record.completed_at = timezone.now()
        record.save()
    except Exception as exc:
        record.status = "failed"
        record.error_message = f"Generation failed: {exc.__class__.__name__}"
        record.completed_at = timezone.now()
        record.save(update_fields=("status", "error_message", "completed_at"))
        raise
    return record
